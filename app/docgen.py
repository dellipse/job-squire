# Copyright (C) 2026 D. Brandmeyer
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.
#
"""Minimal Markdown -> .docx rendering for AI-generated application kits.

This is intentionally not a general Markdown renderer. It handles just the
subset our own kit prompts produce (#/##/### headings, "- "/"* " bullets,
"1. " numbered lists, and **bold** inline emphasis). Anything outside that
subset degrades gracefully to a plain paragraph, so the .docx is always
well-formed even if a model's output drifts from the expected format.
"""
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt

_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
_NUMBERED_RE = re.compile(r"^\d+\.\s+")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, rendering **bold** and *italic* spans as styled runs."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            run = paragraph.add_run(m.group(2))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Render a constrained Markdown subset to a .docx file and return the raw bytes."""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")) and not line.startswith("**"):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
        elif _NUMBERED_RE.match(line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, _NUMBERED_RE.sub("", line))
        elif line.startswith("---") or line.startswith("___") or line.startswith("==="):
            continue  # markdown horizontal rules / separators — skip
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
