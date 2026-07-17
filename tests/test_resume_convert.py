# Copyright (C) 2026 D. Brandmeyer
# Licensed under the GNU Affero General Public License v3 or later.
"""Tests for app/resume_convert.py -- the non-AI document -> markdown
converter used by the "Base Resume" upload path (see
app/main.py:settings_asset_upload)."""
import io

import pytest

from app.resume_convert import (ResumeConversionError, SUPPORTED_EXTENSIONS,
                                convert_to_markdown)


def _docx_bytes(build):
    from docx import Document
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestConvertToMarkdownDispatch:
    def test_unsupported_extension_raises(self):
        with pytest.raises(ResumeConversionError):
            convert_to_markdown(b"whatever", "rtf")

    def test_empty_result_raises(self):
        with pytest.raises(ResumeConversionError):
            convert_to_markdown(b"   \n\n  ", "txt")

    def test_supported_extensions_include_docx_pdf_txt_md(self):
        assert set(SUPPORTED_EXTENSIONS) == {"docx", "pdf", "txt", "md"}

    def test_extension_matching_is_case_and_dot_insensitive(self):
        assert convert_to_markdown(b"hello", "TXT") == "hello"
        assert convert_to_markdown(b"hello", ".txt") == "hello"


class TestTextPassthrough:
    def test_txt_passthrough(self):
        assert convert_to_markdown(b"Jordan Lee\nSummary line.", "txt") == \
            "Jordan Lee\nSummary line."

    def test_md_passthrough(self):
        assert convert_to_markdown(b"# Jordan Lee\n\n- Skill one", "md") == \
            "# Jordan Lee\n\n- Skill one"

    def test_latin1_fallback_decodes(self):
        data = "Café résumé".encode("latin-1")
        assert convert_to_markdown(data, "txt") == "Café résumé"


class TestDocxConversion:
    def test_headings_bold_and_paragraphs(self):
        def build(doc):
            doc.add_heading("Jordan Lee", level=1)
            p = doc.add_paragraph()
            p.add_run("Senior ").bold = True
            p.add_run("Operations Manager")

        markdown = convert_to_markdown(_docx_bytes(build), "docx")
        assert "# Jordan Lee" in markdown
        assert "**Senior** Operations Manager" in markdown

    def test_bullet_list_items_grouped_with_dash_markers(self):
        def build(doc):
            doc.add_paragraph("Experience", style="Heading 2")
            doc.add_paragraph("Led a team of 12", style="List Bullet")
            doc.add_paragraph("Cut costs 20%", style="List Bullet")

        markdown = convert_to_markdown(_docx_bytes(build), "docx")
        assert "- Led a team of 12" in markdown
        assert "- Cut costs 20%" in markdown
        # Consecutive list items land in one block (single blank-line group).
        assert "- Led a team of 12\n- Cut costs 20%" in markdown

    def test_simple_table_converts_to_markdown_table(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Company"
            table.rows[0].cells[1].text = "Role"
            table.rows[1].cells[0].text = "Acme"
            table.rows[1].cells[1].text = "Engineer"

        markdown = convert_to_markdown(_docx_bytes(build), "docx")
        assert "| Company | Role |" in markdown
        assert "| Acme | Engineer |" in markdown

    def test_blank_docx_raises(self):
        def build(doc):
            pass  # no content at all

        with pytest.raises(ResumeConversionError):
            convert_to_markdown(_docx_bytes(build), "docx")

    def test_corrupted_docx_raises_conversion_error_not_crash(self):
        with pytest.raises(ResumeConversionError):
            convert_to_markdown(b"not a real docx file", "docx")


class TestPdfConversion:
    def test_pdf_extracts_text_from_each_page(self, monkeypatch):
        class FakePage:
            def __init__(self, text):
                self._text = text

            def extract_text(self):
                return self._text

        class FakeReader:
            def __init__(self, _stream):
                self.is_encrypted = False
                self.pages = [FakePage("Jordan Lee"), FakePage("Experience section")]

        monkeypatch.setattr("pypdf.PdfReader", FakeReader)
        markdown = convert_to_markdown(b"%PDF-fake", "pdf")
        assert "Jordan Lee" in markdown
        assert "Experience section" in markdown

    def test_encrypted_pdf_raises_conversion_error(self, monkeypatch):
        class FakeReader:
            def __init__(self, _stream):
                self.is_encrypted = True
                self.pages = []

        monkeypatch.setattr("pypdf.PdfReader", FakeReader)
        with pytest.raises(ResumeConversionError):
            convert_to_markdown(b"%PDF-fake", "pdf")

    def test_pdf_with_no_extractable_text_raises(self, monkeypatch):
        class FakePage:
            def extract_text(self):
                return ""

        class FakeReader:
            def __init__(self, _stream):
                self.is_encrypted = False
                self.pages = [FakePage(), FakePage()]

        monkeypatch.setattr("pypdf.PdfReader", FakeReader)
        with pytest.raises(ResumeConversionError):
            convert_to_markdown(b"%PDF-fake", "pdf")
