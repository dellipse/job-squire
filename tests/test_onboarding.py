# Copyright (C) 2026 D. Brandmeyer
# Licensed under the GNU Affero General Public License v3 or later.
"""Tests for the Getting Started walkthrough (``app/onboarding.py``).

The contract (docs/PLAN-onboarding.md, Phase 1): completion is derived from
real app state so the checklist can't drift; steps are skippable and
revisitable; the dashboard card hides on dismissal; account creation is
admin-only and validated; the remote-only board gate respects
``SearchConfig.include_remote``.
"""
import os

import pytest

from app.extensions import db
from app.models import (CandidateAsset, OnboardingState, ProviderCredential,
                        SearchConfig, SearchRun, User)
from tests.conftest import ADMIN_PASSWORD, ADMIN_USERNAME


def _login_admin(client, app):
    from app import _seed_users
    with app.app_context():
        _seed_users(app)
    return client.post("/login",
                       data={"username": ADMIN_USERNAME, "password": ADMIN_PASSWORD},
                       follow_redirects=False)


def _login_user(client, app):
    from app import _seed_users
    from tests.conftest import USER_USERNAME, USER_PASSWORD
    with app.app_context():
        _seed_users(app)
    return client.post("/login",
                       data={"username": USER_USERNAME, "password": USER_PASSWORD},
                       follow_redirects=False)


@pytest.fixture
def clean_state(app_context):
    """Reset onboarding state and the data it derives completion from."""
    def _reset():
        state = db.session.get(OnboardingState, 1)
        if state:
            db.session.delete(state)
        CandidateAsset.query.delete()
        SearchRun.query.delete()
        ProviderCredential.query.delete()
        # test_migrations' mdb fixture drop_all()s the shared DB mid-suite,
        # wiping seeded singletons — recreate rather than assume (same pattern
        # as _login_admin re-seeding accounts).
        cfg = db.session.get(SearchConfig, 1)
        if cfg is None:
            cfg = SearchConfig(id=1)
            db.session.add(cfg)
        cfg.titles, cfg.location = "", ""
        cfg.enabled = False
        db.session.commit()
    _reset()
    from flask import current_app
    profile = os.path.join(current_app.config["DATA_DIR"], "candidate_profile.md")
    if os.path.exists(profile):
        os.unlink(profile)
    yield app_context
    _reset()


class TestChecklistDerivation:
    def test_fresh_install_everything_todo(self, clean_state):
        from app.onboarding import build_checklist
        checklist = build_checklist()
        assert all(item["status"] == "todo" for item in checklist
                   if item["key"] != "accounts")  # seeded 2nd account may exist

    def test_completion_follows_real_data(self, clean_state):
        from app.onboarding import build_checklist, get_state

        state = get_state()
        state.persona = "self"
        db.session.add(CandidateAsset(kind="Resume", original_name="r.pdf",
                                      stored_name="x.pdf"))
        cfg = db.session.get(SearchConfig, 1)
        cfg.titles, cfg.location = "Ops Manager", "Henderson, NV"
        db.session.add(ProviderCredential(provider="themuse", enabled=True))
        db.session.add(SearchRun(status="ok"))
        # Completion requires both the real data AND having visited the step's
        # own page at least once — a step can't complete itself on defaults
        # alone (see OnboardingState.mark_visited).
        for key in ("persona", "profile", "search", "providers", "first_search"):
            state.mark_visited(key)
        db.session.commit()

        status = {i["key"]: i["status"] for i in build_checklist()}
        for key in ("persona", "profile", "search", "providers", "first_search"):
            assert status[key] == "done", key
        assert status["ai"] == "todo"   # nothing AI-ish configured

    def test_data_alone_is_not_enough_without_a_visit(self, clean_state):
        """A step's real data can already satisfy its condition (e.g. the
        seeded "themuse" board is enabled out of the box) without the step
        ever being marked done until its own page has actually been loaded."""
        from app.onboarding import build_checklist
        db.session.add(ProviderCredential(provider="themuse", enabled=True))
        db.session.commit()
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["providers"] == "todo"

    def test_visit_alone_is_not_enough_without_data(self, clean_state, client, app):
        """Loading a step's page shouldn't complete it if nothing was actually
        filled in — both conditions are required."""
        _login_admin(client, app)
        assert client.get("/getting-started/search").status_code == 200
        from app.onboarding import build_checklist
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["search"] == "todo"

    def test_skip_persists_and_revisit_allowed(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.post("/getting-started/ai/skip", follow_redirects=False)
        assert r.status_code == 302
        from app.onboarding import build_checklist
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["ai"] == "skipped"
        # Revisit: the step page still renders and can be completed.
        assert client.get("/getting-started/ai").status_code == 200

    def test_derived_done_beats_stored_skip(self, clean_state):
        from app.onboarding import build_checklist, get_state
        state = get_state()
        state.set_step("search", "skipped")
        state.mark_visited("search")
        cfg = db.session.get(SearchConfig, 1)
        cfg.titles, cfg.location = "Ops Manager", "Henderson, NV"
        db.session.commit()
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["search"] == "done"   # reality wins over the stored skip


class TestDashboardCard:
    def test_card_shows_then_hides_on_dismiss(self, clean_state, client, app):
        _login_admin(client, app)
        # Incomplete + not dismissed: "/" force-redirects into the walkthrough
        # (persona first) rather than rendering the dashboard card.
        r = client.get("/", follow_redirects=False)
        assert r.status_code == 302
        assert "/getting-started/persona" in r.headers["Location"]
        r = client.post("/getting-started/dismiss", follow_redirects=False)
        assert r.status_code == 302
        # Dismissed: "/" renders the plain dashboard again, no more redirect.
        home = client.get("/", follow_redirects=False)
        assert home.status_code == 200
        assert b"Open walkthrough" not in home.data
        # The full page stays reachable from the nav even when dismissed.
        assert client.get("/getting-started").status_code == 200


class TestAccountsStep:
    def test_create_second_account(self, clean_state, client, app):
        _login_admin(client, app)
        User.query.filter(User.role != "admin").delete()
        db.session.commit()
        r = client.post("/getting-started/accounts",
                        data={"username": "jordan", "display_name": "Jordan",
                              "password": "hunter2hunter2", "confirm": "hunter2hunter2"},
                        follow_redirects=False)
        assert r.status_code == 302
        created = User.query.filter_by(username="jordan").first()
        assert created is not None and created.role == "user"
        assert created.check_password("hunter2hunter2")
        db.session.delete(created)
        db.session.commit()

    @pytest.mark.parametrize("data,fragment", [
        ({"username": "x", "password": "hunter2hunter2", "confirm": "hunter2hunter2"},
         "Username"),
        ({"username": "jordan2", "password": "short", "confirm": "short"},
         "at least 8"),
        ({"username": "jordan2", "password": "hunter2hunter2", "confirm": "different"},
         "do not match"),
    ])
    def test_validation_rejects(self, clean_state, client, app, data, fragment):
        _login_admin(client, app)
        r = client.post("/getting-started/accounts", data=data, follow_redirects=True)
        assert fragment.encode() in r.data
        assert User.query.filter_by(username=data["username"]).first() is None

    def test_non_admin_blocked(self, clean_state, client, app):
        _login_user(client, app)
        assert client.get("/getting-started").status_code == 403
        r = client.post("/getting-started/accounts",
                        data={"username": "sneaky", "password": "hunter2hunter2",
                              "confirm": "hunter2hunter2"})
        assert r.status_code == 403
        assert User.query.filter_by(username="sneaky").first() is None


class TestAiStep:
    def test_no_ai_marks_done_and_warns(self, clean_state, client, app):
        _login_admin(client, app)
        assert client.get("/getting-started/ai").status_code == 200  # marks it visited
        r = client.post("/getting-started/ai", data={"action": "no_ai"},
                        follow_redirects=True)
        assert b"Continuing without AI" in r.data
        from app.onboarding import build_checklist
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["ai"] == "done"


class TestNotificationsStep:
    def test_page_renders_both_provider_options(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.get("/getting-started/notifications")
        assert r.status_code == 200
        assert b"SMTP2GO" in r.data
        assert b"apppasswords" in r.data
        assert b"Continue without email notifications" in r.data

    def test_no_email_marks_done_and_warns(self, clean_state, client, app):
        _login_admin(client, app)
        assert client.get("/getting-started/notifications").status_code == 200  # marks visited
        r = client.post("/getting-started/notifications", data={"action": "no_email"},
                        follow_redirects=True)
        assert b"Continuing without email notifications" in r.data
        from app.onboarding import build_checklist
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["notifications"] == "done"

    def test_real_smtp_config_marks_done_once_visited(self, clean_state, client, app):
        """Mirrors the ai step's has_provider check: real, working configuration
        satisfies the step without an explicit "answered"/"no_email" flag, but
        only once the step's own page has been visited (see _step_done)."""
        _login_admin(client, app)
        client.post("/settings/smtp", data={
            "enabled": "on", "host": "mail.smtp2go.com", "port": "587",
            "username": "smtp2go-user", "password": "secret", "from_addr": "me@example.com",
            "to_addr": "me@example.com",
        }, follow_redirects=True)
        from app.onboarding import build_checklist
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["notifications"] == "todo"  # not visited yet
        assert client.get("/getting-started/notifications").status_code == 200
        status = {i["key"]: i["status"] for i in build_checklist()}
        assert status["notifications"] == "done"

    def test_settings_smtp_from_onboarding_returns_to_onboarding(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.post("/settings/smtp", data={
            "enabled": "on", "host": "mail.smtp2go.com", "port": "587",
            "to_addr": "me@example.com", "next": "/getting-started/notifications",
        }, follow_redirects=False)
        assert r.headers["Location"].endswith("/getting-started/notifications")


class TestProfileStep:
    def test_profile_links_append_to_candidate_profile(self, clean_state, client, app):
        from flask import current_app
        _login_admin(client, app)
        path = os.path.join(current_app.config["DATA_DIR"], "candidate_profile.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write("# Existing profile")
        try:
            client.post("/getting-started/profile-links",
                        data={"links": "https://www.linkedin.com/in/jordan\nhttps://github.com/jordan"},
                        follow_redirects=True)
            with open(path, encoding="utf-8") as f:
                text = f.read()
            assert "# Existing profile" in text
            assert "## Online profiles" in text
            assert "- https://www.linkedin.com/in/jordan" in text
        finally:
            os.unlink(path)

    def test_saved_profile_links_render_back_on_page(self, clean_state, client, app):
        """Regression: the profile step must show what's already saved.

        Previously the "Online profiles" textarea was write-only — a saved
        link never appeared anywhere on the page again except in the
        one-time flash message, so a user reloading the step had no way to
        tell it had actually been saved.
        """
        _login_admin(client, app)
        client.post("/getting-started/profile-links",
                    data={"links": "https://www.linkedin.com/in/jordan"},
                    follow_redirects=True)
        r = client.get("/getting-started/profile")
        assert b"https://www.linkedin.com/in/jordan" in r.data

    def test_saved_profile_links_dedupe_across_repeated_saves(self, clean_state, app):
        """save_profile_links() appends a fresh heading each time, so a link
        saved twice must still only be reported once by the read-back helper."""
        from app.onboarding import _saved_profile_links
        with app.app_context():
            from flask import current_app
            path = os.path.join(current_app.config["DATA_DIR"], "candidate_profile.md")
            try:
                with app.test_client() as client:
                    _login_admin(client, app)
                    client.post("/getting-started/profile-links",
                                data={"links": "https://www.linkedin.com/in/jordan"})
                    client.post("/getting-started/profile-links",
                                data={"links": "https://www.linkedin.com/in/jordan\nhttps://github.com/jordan"})
                links = _saved_profile_links()
                assert links == ["https://www.linkedin.com/in/jordan", "https://github.com/jordan"]
            finally:
                if os.path.exists(path):
                    os.unlink(path)


class TestResumeInterview:
    """Phase 2 (docs/PLAN-onboarding.md): the resume-building interview.

    All three transports (manual paste-back, the interactive API chat, and
    the MCP tool) funnel through onboarding.save_resume_draft(), so that
    function carries most of the coverage here.
    """

    def test_save_resume_draft_creates_asset_and_appends_profile(self, clean_state):
        from app.onboarding import save_resume_draft
        result = save_resume_draft(
            "# Jordan Lee\n\nSummary of experience...",
            "Targets operations manager roles, 8 years experience.",
        )
        assert result.get("ok") is True
        asset = CandidateAsset.query.filter_by(kind="Resume").first()
        assert asset is not None
        assert asset.label == "AI-generated resume #1"
        assert asset.is_base is True
        assert asset.id == result["asset_id"]

        from flask import current_app
        profile_path = os.path.join(current_app.config["DATA_DIR"], "candidate_profile.md")
        with open(profile_path, encoding="utf-8") as f:
            profile = f.read()
        assert "From resume interview" in profile
        assert "8 years experience" in profile

    def test_save_resume_draft_creates_new_variant_and_promotes_it(self, clean_state):
        """kind="Resume" is no longer a singleton (docs/PLAN-onboarding.md):
        each save keeps prior drafts around as separate variants and
        promotes only the newest one to is_base."""
        from app.onboarding import save_resume_draft
        first = save_resume_draft("# Draft one")
        second = save_resume_draft("# Draft two, revised")
        assert first["asset_id"] != second["asset_id"]
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 2

        first_asset = db.session.get(CandidateAsset, first["asset_id"])
        second_asset = db.session.get(CandidateAsset, second["asset_id"])
        assert first_asset.is_base is False
        assert second_asset.is_base is True

    def test_save_resume_draft_rejects_blank(self, clean_state):
        from app.onboarding import save_resume_draft
        assert "error" in save_resume_draft("   ")

    def test_save_resume_route_persists(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.post("/getting-started/profile/resume-draft",
                        data={"resume_markdown": "# Test Resume\n\nContent."},
                        follow_redirects=True)
        assert r.status_code == 200
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 1

    def test_resume_interview_non_admin_blocked(self, clean_state, client, app):
        _login_user(client, app)
        assert client.get("/getting-started/profile/interview").status_code == 403

    def test_resume_interview_redirects_without_provider(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.get("/getting-started/profile/interview", follow_redirects=False)
        assert r.status_code == 302
        assert "/getting-started/profile" in r.headers["Location"]

    def test_run_resume_interview_turn_asks_then_completes(self, clean_state, monkeypatch):
        import app.ai as ai_mod

        monkeypatch.setattr(
            ai_mod, "call_with_fallback",
            lambda *a, **k: ("What roles are you targeting?", "test-provider"))
        result = ai_mod.run_resume_interview_turn([])
        assert result["done"] is False
        assert "targeting" in result["message"]

        sentinel_reply = (
            "===RESUME_READY===\n# Jordan Lee\n\nSummary...\n"
            "===PROFILE_FACTS===\nTargets operations manager roles.\n"
        )
        monkeypatch.setattr(
            ai_mod, "call_with_fallback",
            lambda *a, **k: (sentinel_reply, "test-provider"))
        history = [{"role": "assistant", "content": "What roles are you targeting?"},
                  {"role": "user", "content": "Operations manager."}]
        result2 = ai_mod.run_resume_interview_turn(history)
        assert result2["done"] is True
        assert "Jordan Lee" in result2["resume_markdown"]
        assert "operations manager" in result2["profile_facts"].lower()


class TestResumeUploadAutoConvert:
    """Uploading a "Base Resume" document should auto-convert it to markdown
    (app/resume_convert.py) and save it as the kind="Resume" draft, so a
    plain upload satisfies the Getting Started "profile" step the same way
    the resume interview does -- no AI required. See app/main.py:
    settings_asset_upload.
    """

    def _docx_bytes(self):
        import io
        from docx import Document
        doc = Document()
        doc.add_heading("Jordan Lee", level=1)
        doc.add_paragraph("Operations manager with 8 years experience.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_base_resume_docx_upload_creates_resume_draft(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        r = client.post(
            "/settings/assets/upload",
            data={
                "kind": "Base Resume",
                "label": "My resume",
                "file": (io.BytesIO(self._docx_bytes()), "resume.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert r.status_code == 200
        # Both the original upload and the converted markdown draft exist.
        assert CandidateAsset.query.filter_by(kind="Base Resume").count() == 1
        resume_asset = CandidateAsset.query.filter_by(kind="Resume").first()
        assert resume_asset is not None

        from app.onboarding import _read_resume_asset_markdown
        markdown = _read_resume_asset_markdown(resume_asset)
        assert "# Jordan Lee" in markdown
        assert "Operations manager" in markdown

    def test_base_resume_upload_satisfies_profile_step(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        client.post(
            "/settings/assets/upload",
            data={
                "kind": "Base Resume",
                "label": "",
                "file": (io.BytesIO(self._docx_bytes()), "resume.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        # Visiting the profile step marks it visited; the upload above should
        # already satisfy the underlying data condition either way.
        client.get("/getting-started/profile")

        from app.onboarding import _step_data_satisfied, get_state
        state = get_state()
        assert _step_data_satisfied("profile", state) is True

    def test_unsupported_extension_upload_warns_but_does_not_crash(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        r = client.post(
            "/settings/assets/upload",
            data={
                "kind": "Base Resume",
                "label": "",
                "file": (io.BytesIO(b"{\\rtf1 fake rtf}"), "resume.rtf"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert r.status_code == 200
        assert CandidateAsset.query.filter_by(kind="Base Resume").count() == 1
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 0

    def test_non_base_resume_upload_does_not_trigger_conversion(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        client.post(
            "/settings/assets/upload",
            data={
                "kind": "Certification",
                "label": "",
                "file": (io.BytesIO(self._docx_bytes()), "cert.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert CandidateAsset.query.filter_by(kind="Certification").count() == 1
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 0


class TestCustomResumeUploadVariants:
    """Uploading directly under kind="Resume" ("Custom Resume" in the UI) --
    the multi-variant slot itself, not the "Base Resume" archival kind. See
    app/main.py:settings_asset_upload and docs/PLAN-onboarding.md Phase 2.
    """

    def _docx_bytes(self):
        import io
        from docx import Document
        doc = Document()
        doc.add_heading("Jordan Lee", level=1)
        doc.add_paragraph("Operations manager with 8 years experience.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_custom_resume_upload_converts_and_keeps_source(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        r = client.post(
            "/settings/assets/upload",
            data={
                "kind": "Resume",
                "label": "",
                "file": (io.BytesIO(self._docx_bytes()), "resume.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert r.status_code == 200
        # Unlike "Base Resume", nothing else got uploaded -- one row holds
        # both the converted markdown and the original.
        assert CandidateAsset.query.filter_by(kind="Base Resume").count() == 0
        asset = CandidateAsset.query.filter_by(kind="Resume").first()
        assert asset is not None
        assert asset.is_base is True
        assert asset.source_stored_name
        assert asset.source_original_name == "resume.docx"

        from app.onboarding import _read_resume_asset_markdown
        markdown = _read_resume_asset_markdown(asset)
        assert "Jordan Lee" in markdown

    def test_custom_resume_upload_rejects_unconvertible_extension(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        r = client.post(
            "/settings/assets/upload",
            data={
                "kind": "Resume",
                "label": "",
                "file": (io.BytesIO(b"{\\rtf1 fake rtf}"), "resume.rtf"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert r.status_code == 200
        # No raw-binary fallback for this slot -- the upload is rejected
        # outright rather than landing broken content where the profile page
        # assumes markdown (this is the bug that motivated the redesign).
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 0

    def test_second_custom_resume_upload_adds_variant_and_promotes(self, clean_state, client, app):
        import io
        _login_admin(client, app)
        from app.onboarding import save_resume_draft
        first = save_resume_draft("# Draft one")

        client.post(
            "/settings/assets/upload",
            data={
                "kind": "Resume",
                "label": "From another tool",
                "file": (io.BytesIO(self._docx_bytes()), "resume.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=True,
        )
        assert CandidateAsset.query.filter_by(kind="Resume").count() == 2
        first_asset = db.session.get(CandidateAsset, first["asset_id"])
        assert first_asset.is_base is False
        new_asset = (CandidateAsset.query.filter_by(kind="Resume", label="From another tool")
                     .first())
        assert new_asset.is_base is True


class TestResumeSetBaseAndDelete:
    def _make_variant(self, label, is_base=False):
        from flask import current_app
        asset = CandidateAsset(kind="Resume", label=label, original_name=f"{label}.md",
                               stored_name=f"{label}.md", is_base=is_base)
        db.session.add(asset)
        db.session.commit()
        upload_dir = current_app.config["UPLOAD_DIR"]
        with open(os.path.join(upload_dir, asset.stored_name), "w", encoding="utf-8") as f:
            f.write(f"# {label}\n")
        return asset

    def test_set_base_promotes_and_demotes(self, clean_state, client, app):
        _login_admin(client, app)
        a = self._make_variant("A", is_base=True)
        b = self._make_variant("B", is_base=False)

        r = client.post(f"/assets/{b.id}/set-base", data={}, follow_redirects=True)
        assert r.status_code == 200
        db.session.refresh(a)
        db.session.refresh(b)
        assert a.is_base is False
        assert b.is_base is True

    def test_set_base_rejects_non_resume_kind(self, clean_state, client, app):
        _login_admin(client, app)
        cert = CandidateAsset(kind="Certification", original_name="c.pdf", stored_name="c.pdf")
        db.session.add(cert)
        db.session.commit()
        r = client.post(f"/assets/{cert.id}/set-base", data={}, follow_redirects=True)
        assert r.status_code == 200
        db.session.refresh(cert)
        assert cert.is_base is False

    def test_deleting_base_variant_promotes_newest_remaining(self, clean_state, client, app):
        _login_admin(client, app)
        a = self._make_variant("A", is_base=True)
        b = self._make_variant("B", is_base=False)

        r = client.post(f"/assets/{a.id}/delete", data={}, follow_redirects=True)
        assert r.status_code == 200
        assert db.session.get(CandidateAsset, a.id) is None
        db.session.refresh(b)
        assert b.is_base is True


class TestRemoteGate:
    def _run(self, monkeypatch, include_remote):
        import app.search as search_mod
        cfg = db.session.get(SearchConfig, 1)
        cfg.titles, cfg.location = "Ops Manager", "Henderson, NV"
        cfg.enabled = True
        cfg.include_remote = include_remote
        db.session.add(ProviderCredential(provider="jobicy", enabled=True))
        db.session.add(ProviderCredential(provider="themuse", enabled=True))
        db.session.commit()
        called = []

        def fake_search(provider, creds, titles, cfg_dict):
            called.append(provider)
            return [], None

        monkeypatch.setattr(search_mod, "search_provider", fake_search)
        monkeypatch.setattr(search_mod, "_maybe_email_digest",
                            lambda *a, **k: None, raising=False)
        search_mod.run_search(trigger="manual")
        return called

    def test_remote_only_board_skipped_when_off(self, clean_state, monkeypatch):
        called = self._run(monkeypatch, include_remote=False)
        assert "jobicy" not in called
        assert "themuse" in called

    def test_remote_only_board_runs_when_on(self, clean_state, monkeypatch):
        called = self._run(monkeypatch, include_remote=True)
        assert "jobicy" in called


class TestSafeNext:
    def test_relative_honored_absolute_rejected(self, clean_state, client, app):
        _login_admin(client, app)
        r = client.post("/settings/search",
                        data={"titles": "Ops", "location": "Henderson, NV",
                              "country": "US", "next": "/getting-started/search"},
                        follow_redirects=False)
        assert r.headers["Location"].endswith("/getting-started/search")
        r = client.post("/settings/search",
                        data={"titles": "Ops", "location": "Henderson, NV",
                              "country": "US", "next": "https://evil.example/x"},
                        follow_redirects=False)
        assert "evil.example" not in r.headers["Location"]
